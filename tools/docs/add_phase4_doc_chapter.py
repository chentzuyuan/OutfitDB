#!/usr/bin/env python3
"""Phase 4 chapter — clean append-to-end with proper Heading styles.
First strips out the previous failed insertion (paragraphs 521-611), then
appends Chapter 18 + restores Appendix A/B at the end."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

DOC_PATH = Path("/Users/buttegg/Desktop/ClosetMind_Documentation.docx")

doc = Document(str(DOC_PATH))

# Find the previous (failed) insertion: it starts at "18. Phase 4: ..."
# and runs until just before "Appendix A".
phase4_start = None
appendix_a_start = None
for i, p in enumerate(doc.paragraphs):
    if "18. Phase 4: Multi-Stage" in (p.text or ""):
        phase4_start = i
    if (p.text or "").strip().startswith("Appendix A") and appendix_a_start is None:
        appendix_a_start = i

# Snapshot Appendix A/B content (paragraphs from appendix_a_start onward)
# so we can re-emit them after rewriting Ch 18.
appendix_paras = doc.paragraphs[appendix_a_start:]
appendix_data = []
for p in appendix_paras:
    style = p.style.name if p.style else "Normal"
    appendix_data.append((p.text, style))

# Remove the failed Ch 18 + Appendix paragraphs (delete from phase4_start onward).
body = doc.paragraphs[0]._parent
for p in list(doc.paragraphs[phase4_start:]):
    p._element.getparent().remove(p._element)

# ─── Rewrite Chapter 18 with proper styles ───────────────────────────────
# The doc has duplicate "Heading 1/2/3" styles, breaking the name lookup.
# Resolve by enumerating styles and grabbing the LAST occurrence (custom one).
heading_styles = {}
for s in doc.styles:
    if s.name in ("Heading 1", "Heading 2", "Heading 3"):
        heading_styles[s.name] = s

def add_h(text, level):
    p = doc.add_paragraph(text)
    style_name = f"Heading {level}"
    if style_name in heading_styles:
        p.style = heading_styles[style_name]
    return p

def add_p(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    return p

def add_code(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Menlo"
        run.font.size = Pt(9)

def add_bullet(text):
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        p = doc.add_paragraph("• " + text)
        p.paragraph_format.left_indent = Inches(0.25)
    return p


add_h("18. Phase 4: Multi-Stage Training Architecture", 1)
add_p(
    "Phase 4 separates the recommendation problem into three independent learning tracks "
    "so user feedback isn't polluted by off-axis confounds. The original calibration mixed "
    "temperature, aesthetic, and occasion judgments into a single rating, which caused the "
    "model to learn cross-contaminated signals (e.g., users would down-rate a "
    "well-designed outfit because the temperature was wrong, or vice versa). The Phase 4 "
    "architecture asks the user three separate questions in sequence and trains three "
    "separate XGBoost classifiers, then chains them at inference time."
)

add_h("18.1 Motivation", 2)
add_p(
    "Three issues drove the redesign. First, observed outfits like 'Down Puffer + Cargo Shorts "
    "at 14°C formal' were getting positive ratings from the SHAP-driven loop because their "
    "individual items had high historical averages, even though no human would wear that combo. "
    "Second, the single-rating UI made it impossible for the user to convey 'great look but wrong "
    "temperature.' Third, hand-coded scoring rules (heat ceiling, occasion formality) worked but "
    "couldn't be personalized — every user got the same rules."
)
add_p(
    "Phase 4's contract is: the system asks pure-axis questions, the user gives pure-axis "
    "answers, three models learn pure-axis decisions. At inference time, the three signals are "
    "multiplied together so any single hard fail (wrong temperature, wrong occasion) drives the "
    "score to zero regardless of aesthetic preference."
)

add_h("18.2 The Three Stages", 2)
add_h("18.2.1 Stage 1 — Temperature", 3)
add_p(
    "User sees an outfit and checks the temperature zones it would suit. Multi-select; empty "
    "selection means 'not appropriate for any zone.' Six zones replace the previous five-zone "
    "calibration to add a sub-zero category for snow / freezing scenarios:"
)
add_bullet("subzero — < 0°C (snow, ice)")
add_bullet("cold — 0–10°C")
add_bullet("cool — 10–18°C")
add_bullet("mild — 18–25°C")
add_bullet("warm — 25–30°C")
add_bullet("hot — > 30°C")
add_p(
    "The candidate generator for Stage 1 produces 80% deliberately-random combinations "
    "(tank top + heavy coat + cargo shorts) and 20% baseline cohesive outfits. The random "
    "combos are critical: they let the model learn what 'too warm for this zone' looks like, "
    "not just what 'cohesive cool-zone outfits' look like."
)

add_h("18.2.2 Stage 2 — Aesthetic", 3)
add_p(
    "User rates outfits on a four-point scale: dislike (-1), meh (0), like (1), love (2). "
    "The candidate generator uses the existing cohesive-outfit pipeline so the user is never "
    "shown obviously-broken combinations — they are judging color/cut/layering quality only, "
    "with temperature and occasion held constant by virtue of the candidate pool. This stage "
    "reuses the existing Rating table and XGBoost training pipeline from Phase 1."
)

add_h("18.2.3 Stage 3 — Occasion", 3)
add_p(
    "User checks which events an outfit would fit. Nine events are defined as a fixed standard "
    "so the model learns over a controlled vocabulary:"
)
events = [
    "home — at home / lounging",
    "gym — workout / sport",
    "beach — vacation / coastal",
    "casual_outing — friend hangout / shopping",
    "date_night — dinner date",
    "interview — job interview",
    "office — daily work",
    "business_meeting — client meeting / presentation",
    "formal_event — wedding / gala / formal dinner",
]
for txt in events:
    add_bullet(txt)
add_p(
    "Each event has an associated formality level (0–5) used as a fallback before the model "
    "is trained. Once trained, the learned model overrides these fallback rules with the "
    "user's personal interpretation of what 'office-appropriate' means in their wardrobe."
)

add_h("18.3 Adaptive Sampling for Stage 1", 2)
add_p(
    "Stage 1 uses adaptive sampling so coverage doesn't depend on the user clicking through "
    "every zone equally. After each batch the system inspects per-zone statistics and picks "
    "the next target zone using two rules:"
)
add_bullet(
    "Coverage gap: if any zone has fewer than MIN_PER_ZONE (15) ratings, prioritize sampling "
    "for that zone."
)
add_bullet(
    "Drift detection: once all zones meet MIN_PER_ZONE, check acceptance rates. Any zone with "
    "less than 30% acceptance triggers an additional probe batch — the system pushes outfit "
    "warmth toward whatever the user has previously accepted in that zone, narrowing on their "
    "personal comfort range."
)
add_p(
    "The practical effect: a cold-sensitive user who rejects every cool-zone outfit at "
    "moderate warmth will receive progressively-warmer cool-zone outfits until the system "
    "finds their threshold. A heat-sensitive user gets the opposite. Stage 1 completes when "
    "all six zones are covered AND no zone shows a drift signal — typically 90 to 150 ratings, "
    "depending on how unusual the user's preferences are."
)

add_h("18.4 Schema Additions", 2)
add_p("Phase 4 adds two tables and four user-level columns:")
add_p("temperature_ratings", bold=True)
add_bullet("id, user_id, outfit_id (unique constraint on user×outfit)")
add_bullet("zones_ok: JSON list — subset of TEMP_ZONE_KEYS (may be empty)")
add_bullet("target_zone: which zone we sampled this outfit for (analytics)")
add_p("occasion_ratings", bold=True)
add_bullet("id, user_id, outfit_id")
add_bullet("events_ok: JSON list — subset of EVENT_KEYS")
add_p("users (new columns)", bold=True)
add_bullet("v2_temp_done, v2_aesthetic_done, v2_occasion_done — booleans")
add_bullet("zone_warmth_prefs — JSON dict {zone: [min_warmth, max_warmth]} for "
           "future per-user warmth profile beyond a single offset")

add_h("18.5 Training Pipelines", 2)
add_p(
    "Stages 1 and 3 frame multi-label classification as binary classification with the "
    "target label one-hot-encoded into the feature vector. So one model handles all six "
    "temperature zones (Stage 1) or all nine events (Stage 3), and inference asks 'is this "
    "outfit OK for THIS zone/event?' rather than 'which zones is this outfit OK for?'."
)
add_p("Concretely, each TemperatureRating row expands into six training samples:")
add_code([
    "for zk in TEMP_ZONE_KEYS:",
    "    features = build_feature_vector(outfit, ctx)  # 100-dim",
    "    one_hot  = zone_one_hot(zk)                   # 6-dim",
    "    label    = 1 if zk in zones_ok else 0",
    "    X.append(concat(features, one_hot))           # 106-dim",
    "    y.append(label)",
])
add_p(
    "Stage 3 mirrors this with 9-dimensional event one-hot. Both stages use XGBoost "
    "(n_estimators=200, max_depth=4, learning_rate=0.05, scale_pos_weight balanced for "
    "class imbalance, early stopping at 20 rounds). Trained models live at "
    "models/stage1_temp.json and models/stage3_occasion.json. Stage 2 keeps using the "
    "existing models/current.json from Phase 1's model_training.train_model()."
)
add_p(
    "Auto-retrain triggers on every batch submit. Each /training/v2/*/submit endpoint "
    "calls the corresponding training function as part of its response, so the model "
    "evolves continuously as the user goes through the flow."
)

add_h("18.6 Scoring Chain", 2)
add_p(
    "Final score combines all three stages multiplicatively, so a hard fail in any stage "
    "drives the total to near zero:"
)
add_code([
    "aesthetic_blend = 0.65*pref + 0.20*fresh + 0.10*div + 0.05*recovery",
    "total = ctx_fit  ×  stage1_temp_pass  ×  stage3_occasion_pass  ×  aesthetic_blend",
])
add_p("Each multiplier's role:")
add_bullet(
    "ctx_fit (0.0–1.0): legacy soft-temperature score from the Layer Coverage Model — "
    "warmth_score versus ideal_warmth_for_temp. Always available."
)
add_bullet(
    "stage1_temp_pass (0.0–1.0): learned probability the outfit suits the current temp zone. "
    "Cold-start fallback: 1.0 (no filtering) until the model is trained."
)
add_bullet(
    "stage3_occasion_pass (0.0–1.0): learned probability the outfit suits the current event. "
    "Cold-start fallback: 1.0."
)
add_bullet(
    "aesthetic_blend: weighted blend of preference (Stage 2 XGBoost), freshness, diversity, "
    "and recovery."
)
add_p(
    "The chain is read by humans as: 'Is the temperature appropriate? AND is the occasion "
    "appropriate? AND does the user like how it looks?' — three independent gates, each "
    "learnable from its own pure-axis training data."
)

add_h("18.7 Hard Rules vs Learned Rules", 2)
add_p(
    "Phase 4 keeps three rule-based filters at the candidate-generation stage. These are "
    "not redundant with the learned models — they serve as guardrails that prevent the "
    "candidate pool from ever containing universally-bad combinations, regardless of what "
    "any user has labeled."
)
add_p("Rule 1: Temperature gate (outfit_generator._passes_temperature_gate)", bold=True)
add_p(
    "Outfit warmth must satisfy 0.6 × ideal_warmth ≤ warmth ≤ 1.5 × ideal_warmth. "
    "Below the lower bound is the underdressed case; above the upper bound is the "
    "overdressed case that earlier rule sets missed."
)
add_p("Rule 2: Occasion-formality gate (outfit_generator._passes_occasion_gate)", bold=True)
add_p(
    "Average core-item formality (tops + bottom + shoes + fullbody) must meet the occasion's "
    "minimum required formality. Additionally, no single core item may be more than 1.5 "
    "levels below the requirement (the weak-link check). This catches outfits like "
    "'tuxedo trousers + tee + sneakers' that pass the average check but fail the dress code "
    "via a single visible item."
)
add_p("Rule 3: Outer-alone gate (outfit_generator._legal)", bold=True)
add_p(
    "If the only top items are outer-role (jackets, coats), at least one of them must have "
    "Item.can_wear_alone = True. This rule is data-driven, not gender-coded: typical "
    "menswear outerwear (peacoat, wool overcoat, leather jacket) ships with can_wear_alone=False, "
    "while knitwear and statement pieces (cardigans, hoodies, kimonos) ship with True. "
    "Future women's-wear or androgynous profiles can flip individual items without code changes."
)
add_p(
    "The relationship: hard rules carve out the search space; learned models rank within it. "
    "Hard rules can never override the user's per-item training preferences — they only set "
    "the floor of 'physically/logically valid outfits.'"
)

add_h("18.8 UI Flow", 2)
add_p("Three sub-pages live under /training:")
add_bullet(
    "/training/temperature — six-checkbox UI, multi-select per outfit, adaptive zone targeting "
    "shown in the header, per-zone stats panel."
)
add_bullet(
    "/training/aesthetic — four-button rating UI (Dislike / Meh / Like / Love), one rating per "
    "outfit, total-progress bar."
)
add_bullet(
    "/training/occasion — nine-checkbox UI, multi-select per outfit, total-progress bar."
)
add_p(
    "Sequence is recommended (temperature → aesthetic → occasion) but not enforced. "
    "Each page is self-contained and writes only to its own ratings table, so users can "
    "interleave or revisit any stage without invalidating the others."
)

add_h("18.9 Observed Behavior on the Tester Profile", 2)
add_p(
    "Before Phase 4, the Tester profile (54 items, ~2700 ratings) showed roughly "
    "8/10 obviously-bad top-1 recommendations across the canonical scenario suite. After "
    "the hard rule additions and the three-stage chain, the same scenarios produce "
    "4/10 strong recommendations, 5/10 with minor stylistic quirks, and 1/10 wardrobe "
    "gap (28°C sport — the wardrobe lacks pure athletic tops). The remaining stylistic "
    "quirks (trench coat at 24°C casual borderline-warm, tuxedo trousers at casual being "
    "over-formal) are exactly the cases the learned models will personalize over time as "
    "the user labels them with stages 1 and 3."
)

# ─── Re-emit Appendix A/B at the end ─────────────────────────────────────
for txt, style_name in appendix_data:
    if style_name in ("Heading 1", "Heading 2", "Heading 3"):
        p = doc.add_paragraph(txt)
        if style_name in heading_styles:
            p.style = heading_styles[style_name]
    else:
        doc.add_paragraph(txt)

doc.save(str(DOC_PATH))
print(f"Saved with proper styles. Total paragraphs: {len(doc.paragraphs)}")
